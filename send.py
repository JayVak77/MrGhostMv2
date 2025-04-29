import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
import pyfiglet
from html2docx import html2docx
from pdfkit import from_string
from io import BytesIO
import os
import json
import random
import time
from datetime import datetime
import signal
import sys
import concurrent.futures
from exchangelib import DELEGATE, Account, Credentials, Message, HTMLBody, Configuration, Mailbox, FileAttachment, NTLM, OAuth2Credentials
from exchangelib.protocol import Protocol
import string
import hashlib
from datetime import datetime, timedelta
import base64
import re
import requests
from faker import Faker
from email.utils import formataddr
from email.mime.base import MIMEBase
from email import encoders
import dns.resolver
import socks
import socket
import email
from email import policy
from email.parser import BytesParser
import gnupg
from email import message_from_string
from email.utils import parseaddr
import subprocess
import pdfkit
import tempfile
import uuid
from docx import Document
from docx.shared import Inches
from bs4 import BeautifulSoup
import io
import qrcode
import quopri  # Correct import
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from art import text2art
import threading
import platform
import os
import sys
import json

# --- base_path function ---
def get_base_path():
    if hasattr(sys, '_MEIPASS'):
        return sys._MEIPASS
    else:
        return os.path.dirname(os.path.abspath(__file__))


# --- Load configuration ---
base_path = get_base_path()
config_path = os.path.join(base_path, 'config.json')
with open(config_path, 'r', encoding='utf-8') as config_file:
    config = json.load(config_file)

# Load other files similarly...

# Licensing
# Define the API URLs
LICENSE_API_URL = "https://tools.thrilldigitals.com/api/validate-license"

LICENSE_FILE = "license.txt"
STATIC_PRODUCT_ID = 1

def get_pc_identifier():
    # Generate a unique identifier for the PC
    unique_str = platform.node() + platform.processor() + platform.machine()
    return hashlib.sha256(unique_str.encode()).hexdigest()

def validate_license(license_key, pc_identifier):
    # Send a POST request with license key, PC identifier, and static product ID
    response = requests.post(LICENSE_API_URL, data={
        'license_key': license_key,
        'pc_identifier': pc_identifier,
        'product_id': STATIC_PRODUCT_ID  # Include static product_id
    })

    if response.status_code == 200 and response.json().get('success'):
        print("License validated successfully.")
        # Save the license key to a file
        with open(LICENSE_FILE, 'w') as f:
            f.write(license_key)
        return True
    else:
        print("License validation failed:", response.json().get('error'))
        return False

def load_license():
    # Check if the license file exists and read the key from it
    if os.path.exists(LICENSE_FILE):
        with open(LICENSE_FILE, 'r') as f:
            return f.read().strip()
    return None

def prompt_license_key(pc_identifier):
    license_key = input("Please enter your license key: ")
    if not validate_license(license_key, pc_identifier):
        print("Application cannot be started due to invalid license.")
        exit(1)
    return license_key

def check_license_continuously(license_key, pc_identifier):
    while True:
        if not validate_license(license_key, pc_identifier):
            print("License validation failed during usage. Please enter a valid license key.")
            license_key = prompt_license_key(pc_identifier)
        # Check the license every 60 seconds
        time.sleep(60)

#End Licensing

# Get the directory of the current script
script_dir = os.path.dirname(os.path.abspath(__file__))

# Signal handler for graceful shutdown
should_stop = False
def signal_handler(signum, frame):
    global should_stop
    should_stop = True
    print("\n\033[93m[INFO] Stopping the script. Please wait for the current email to finish sending...\033[0m")
    sys.exit(0)

signal.signal(signal.SIGINT, signal_handler)

def print_banner():
    banner = text2art("MrGhost - Mailer")
    print(banner)

def print_centered(text, color_code="\033[92m"):
    terminal_width = os.get_terminal_size().columns
    centered_text = text.center(terminal_width)
    print(f"{color_code}{centered_text}\033[0m")

def print_summary(successful_sends, failed_sends, total_emails):
    print("\n")
    summary_text = f"Summary: Sent: {successful_sends} | Failed: {failed_sends} | Total: {total_emails}"
    print_centered(summary_text)
    if config.get("use_single_smtp", False):
        smtp_host, smtp_port, smtp_user, _ = config['smtp_server'].split('|')
        sender_info = f"Sender: {config['single_smtp_sender_email']} | SMTP: {smtp_host}:{smtp_port}"

    sender_info = "MrGhost Tools and Services"
    print_centered(sender_info, "\033[94m")
    print_centered("=" * 20 + " MrGhost Mailer V 1.2 " + "=" * 20, "\033[93m")

def generate_random_values():
    return {
        'random_number10': ''.join(random.choices(string.digits, k=10)),
        'random_phonenumber': f"+1{''.join(random.choices(string.digits, k=10))}",
        'random_string': ''.join(random.choices(string.ascii_letters + string.digits, k=10)),
        'random_md5': hashlib.md5(random.randbytes(10)).hexdigest(),
        'random_path': ''.join(random.choices(string.ascii_lowercase, k=5)),
    }


def generate_qr_code(url, recipient_email, random_values, box_size=10, border=4):
    qr = qrcode.QRCode(version=1, box_size=box_size, border=border)
    qr.add_data(url)
    qr.make(fit=True)
    img = qr.make_image(fill_color=config['qrcode_color'], back_color=config['qrcode_bg_color'])
    buffered = io.BytesIO()
    img.save(buffered, format="PNG")
    img_str = base64.b64encode(buffered.getvalue()).decode()
    return f"data:image/png;base64,{img_str}"

def send_email_via_exchangelib(to_email, msg, sender_name, exchange_email, exchange_server, exchange_password):
    try:
        # Create credentials and configuration
        credentials = Credentials(username=exchange_email, password=exchange_password)
        config = Configuration(server=exchange_server, credentials=credentials)

        # Create account
        account = Account(
            primary_smtp_address=exchange_email,
            config=config,
            autodiscover=False,
            access_type=DELEGATE
        )

        # Extract HTML content from the message
        html_content = None
        for part in msg.walk():
            if part.get_content_type() == "text/html":
                html_content = part.get_payload(decode=True).decode(part.get_content_charset())
                break

        if not html_content:
            raise ValueError("No HTML content found in the email message")

        # Create and configure the message
        m = Message(account=account)
        m.subject = msg['Subject']
        m.body = HTMLBody(html_content)
        m.to_recipients = [Mailbox(email_address=to_email)]
        m.sender = Mailbox(email_address=exchange_email, name=sender_name)

        # Handle multiple inline images using relative paths
        format_dir = resource_path('format')
        image_paths = {
            'a1': os.path.join(format_dir, 'a1.png'),
            'a2': os.path.join(format_dir, 'a2.png'),
            'a3': os.path.join(format_dir, 'a3.png')
        }

        # Attach all images
        for cid, image_path in image_paths.items():
            if os.path.exists(image_path):
                with open(image_path, "rb") as image_file:
                    file_content = image_file.read()

                image_attachment = FileAttachment(
                    name=f"{cid}.png",
                    content=file_content,
                    is_inline=True,
                    content_id=cid
                )
                m.attach(image_attachment)
            else:
                print(f"\033[91m[ERROR] Image file not found at path: {image_path}\033[0m")

        # Handle other attachments
        for part in msg.walk():
            if part.get_content_maintype() == 'application' and part.get_filename() not in [f"{cid}.png" for cid in image_paths.keys()]:
                file = FileAttachment(
                    name=part.get_filename(),
                    content=part.get_payload(decode=True)
                )
                m.attach(file)

        # Send the email
        m.send(save_copy=False)
        return True
    except Exception as e:
        print(f"\033[91m[ERROR] Failed to send email via Exchange to {to_email}: {str(e)}\033[0m")
        return False


def send_email_via_owa(to_email, subject, body, sender_name, fake_company_email, owa_config):
    try:
        print(f"\033[94m[DEBUG] OWA Config received: {owa_config}\033[0m")
        print(f"\033[94m[INFO] Attempting to send email via OWA to {to_email}\033[0m")

        if 'owa_server' not in owa_config:
            raise KeyError("'owa_server' not found in owa_config")

        owa_server, owa_email, owa_password = owa_config['owa_server'].split('|')

        credentials = Credentials(username=owa_email, password=owa_password)
        protocol = Protocol(
            type=NTLM,
            server=owa_server,
            verify_ssl=owa_config.get('owa_use_ssl', True)
        )

        account = Account(
            primary_smtp_address=owa_email,
            credentials=credentials,
            autodiscover=False,
            access_type=DELEGATE,
            protocol=protocol
        )

        message = Message(
            account=account,
            subject=subject,
            body=HTMLBody(body),
            to_recipients=[Mailbox(email_address=to_email)],
            sender=Mailbox(email_address=fake_company_email, name=sender_name)
        )

        # Handle attachment if needed
        if owa_config.get('send_attachment', False):
            attachment_path = owa_config.get('attachment_path', '')
            if attachment_path and os.path.exists(attachment_path):
                with open(attachment_path, 'rb') as file:
                    content = file.read()
                    attachment = FileAttachment(name=os.path.basename(attachment_path), content=content)
                    message.attach(attachment)

        message.send_and_save()
        print(f"\033[95m[SUCCESS] Email sent to {to_email} via OWA\033[0m")
        return True
    except Exception as e:
        print(f"\033[91m[ERROR] Failed to send email via OWA to {to_email}: {e}\033[0m")
        return False

# def send_emails_concurrently(email_list, body):
#     successful_sends = 0
#     failed_sends = 0
#     total_emails = len(email_list)

#     if total_emails == 0:
#         print("\033[91m[ERROR] The email list is empty. No emails to send.\033[0m")
#         return 0, 0

#     try:
#         for count, email in enumerate(email_list, 1):
#             email = email.strip()
#             if not email:  # Skip empty lines in the email list
#                 continue

#             success = send_email(email, body, count)
#             if success:
#                 successful_sends += 1
#                 print(f"\033[95m[{count}/{total_emails}] Email sent successfully to {email}\033[0m")
#             else:
#                 failed_sends += 1
#                 print(f"\033[91m[{count}/{total_emails}] Failed: {email}\033[0m")

#             # Pause after sending a certain number of emails
#             if count % config['smtp_pause_num'] == 0:
#                 print(f"\033[94m[INFO] Pausing after sending {config['smtp_pause_num']} emails.\033[0m")
#                 time.sleep(config['smtp_sleep_time'])

#         return successful_sends, failed_sends
#     except Exception as e:
#         print(f"\033[91m[ERROR] Error in concurrent sending: {str(e)}\033[0m")
#         return 0, 0

def send_emails_concurrently(email_list, body):
    successful_sends = 0
    failed_sends = 0
    pause_count = 0
    total_emails = len(email_list)

    # Display attachment information once at the beginning
    if config.get('send_attachment', False):
        attachment_path = config.get('attachment_path', '')
        if attachment_path:
            full_attachment_path = resource_path(attachment_path)
        if os.path.exists(full_attachment_path):
            print(f"\033[94m[INFO] Attachment will be added: {attachment_path}\033[0m")
        else:
            print(f"\033[91m[WARNING] Attachment file not found: {attachment_path}\033[0m")

    try:
        with concurrent.futures.ThreadPoolExecutor(max_workers=5) as executor:
            futures = {executor.submit(send_email, email.strip(), body, count): count
                      for count, email in enumerate(email_list, 1)}

            for future in concurrent.futures.as_completed(futures):
                count = futures[future]
                email = email_list[count - 1].strip()

                if should_stop:
                    print("\033[93m[INFO] Script stopped by user.\033[0m")
                    break

                try:
                    if future.result():
                        successful_sends += 1
                        masked_email = mask_email(email)
                        print(f"\033[95m[{count}/{total_emails}] Email sent successfully to {masked_email} via Exchange\033[0m")
                    else:
                        failed_sends += 1
                        masked_email = mask_email(email)
                        print(f"\033[91m[{count}/{total_emails}] Failed: {masked_email}\033[0m")
                except Exception as e:
                    masked_email = mask_email(email)
                    print(f"\033[91m[ERROR] {masked_email}: {e}\033[0m")
                    failed_sends += 1

                # Pause after sending a certain number of emails
                if count % config['smtp_pause_num'] == 0:
                    pause_count += 1
                    print(f"\033[94m[INFO] Pausing after sending {config['smtp_pause_num']} emails. Pause #{pause_count}\033[0m")
                    time.sleep(config['smtp_sleep_time'])

        return successful_sends, failed_sends
    except Exception as e:
        print(f"\033[91m[ERROR] Error in concurrent sending: {str(e)}\033[0m")
        return 0, 0  # Return tuple of zeros in case of error


def send_email_via_smtp(to_email, msg, smtp_config):
    try:
        with smtplib.SMTP(smtp_config['host'], smtp_config['port'], timeout=30) as server:
            if smtp_config.get('use_tls', False):
                server.starttls()
            server.login(smtp_config['username'], smtp_config['password'])
            server.sendmail(smtp_config['username'], to_email, msg.as_string())
        return True
    except Exception as e:
        print(f"\033[91m[ERROR] Failed to send email via SMTP: {str(e)}\033[0m")
        return False


def send_email(to_email, body, count):
    if should_stop:
        return False

    random_values = generate_random_values()

    # Load and prepare email content from config.json template path
    html_template_path = os.path.join(base_path, config['html_template_path'])
    with open(html_template_path, 'r', encoding='utf-8') as html_file:
        html_content = html_file.read()

    # Placeholder replacements
    email_body = replace_placeholders(html_content, to_email, random_values)
    sender_name = random.choice(config['sender_names']) if config.get('sender_names') else "Default Sender"
    # subject = random.choice(config['subjects'])
    # subject = replace_placeholders(subject, to_email, random_values)
    subject = replace_placeholders(random.choice(config['subjects']), to_email, random_values)


    # Create MIME message
    msg = MIMEMultipart('mixed')
    msg['Subject'] = subject
    msg['From'] = formataddr((sender_name, config['single_smtp_sender_email']))
    msg['To'] = to_email
    msg['Date'] = time.strftime("%a, %d %b %Y %H:%M:%S +0000", time.gmtime())
    msg.attach(MIMEText(email_body, 'html'))

    # Handle attachments if configured
    if config.get('send_attachment', False):
        attachment_path = os.path.join(base_path, config['attachment_path'])
        if os.path.exists(attachment_path):
            with open(attachment_path, 'rb') as attachment_file:
                attachment = MIMEBase('application', 'octet-stream')
                attachment.set_payload(attachment_file.read())
                encoders.encode_base64(attachment)
                attachment.add_header('Content-Disposition', 'attachment', filename=os.path.basename(attachment_path))
                msg.attach(attachment)
    else:
        print(f"Attachment not found at {attachment_path}")

    # Determine the sending method based on the config
    if config.get('use_exchangelib', False):
        # Directly retrieve these parameters instead of splitting the 'exchange_server' string
        exchange_email = config.get('exchange_email')  # Get the email address
        exchange_password = config.get('exchange_password')  # Get the password
        exchange_server = config.get('exchange_server')  # Get the server address

        if not exchange_email or not exchange_password or not exchange_server:
            print(f"\033[91m[ERROR] Invalid Exchange server configuration: Missing email, password or server.\033[0m")
            return False

        return send_email_via_exchangelib(to_email, msg, sender_name, exchange_email, exchange_server, exchange_password)

    elif config.get('single_smtp', False):
        # Send via single SMTP server
        smtp_config = {
            'host': config['smtp_server'].split('|')[0],
            'port': int(config['smtp_server'].split('|')[3]),
            'username': config['smtp_server'].split('|')[1],
            'password': config['smtp_server'].split('|')[2],
            'use_tls': config['use_tls']
        }
        return send_email_via_smtp(to_email, msg, smtp_config)

    elif config.get('smtp_servers', ''):
        # Send emails via multiple SMTP servers listed in smtp.txt
        with open(config['smtp_servers'], 'r') as file:
            smtp_server_list = file.readlines()

        successful_sends, failed_sends = 0, 0
        for smtp_server in smtp_server_list:
            smtp_server = smtp_server.strip()
            if len(smtp_server.split('|')) < 5:
                print(f"\033[91m[ERROR] Invalid SMTP server configuration: {smtp_server}\033[0m")
                continue
            smtp_config = {
                'host': smtp_server.split('|')[0],
                'port': int(smtp_server.split('|')[3]),
                'username': smtp_server.split('|')[1],
                'password': smtp_server.split('|')[2],
                'use_tls': config['use_tls']
            }

            success = send_email_via_smtp(to_email, msg, smtp_config)
            if success:
                successful_sends += 1
            else:
                failed_sends += 1

        return successful_sends, failed_sends

    else:
        print("\033[91m[ERROR] No valid email sending method specified in config.json.\033[0m")
        return False


def html_to_word(html_content, recipient_email, random_values):
    try:
        # Replace placeholders and generate QR code
        html_content = replace_placeholders(html_content, recipient_email, random_values)
        if '{QR_CODE_PLACEHOLDER}' in html_content:
            qr_code_img = generate_qr_code(random.choice(config['links']), recipient_email, random_values)
            html_content = html_content.replace('{QR_CODE_PLACEHOLDER}', f'<img src="{qr_code_img}" alt="QR Code">')

        soup = BeautifulSoup(html_content, 'html.parser')
        doc = Document()

        for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'ul', 'ol', 'img']):
            if element.name == 'img' and element.get('src', '').startswith('data:image/png;base64,'):
                image_data = base64.b64decode(element['src'].split(',')[1])
                image_stream = io.BytesIO(image_data)
                doc.add_picture(image_stream, width=Inches(2))
            elif element.name == 'p':
                doc.add_paragraph(element.text)
            elif element.name.startswith('h'):
                doc.add_heading(element.text, level=int(element.name[1]))
            elif element.name in ['ul', 'ol']:
                for li in element.find_all('li'):
                    doc.add_paragraph(li.text, style='List Bullet' if element.name == 'ul' else 'List Number')

        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_content = docx_buffer.getvalue()

        return docx_content
    except Exception as e:
        print(f"\033[91m[ERROR] Failed to convert HTML to DOCX: {str(e)}\033[0m")
        return None

def html_to_pdf(html_content, recipient_email, random_values, image_paths=None):
    try:
        # Define the path to wkhtmltopdf for Windows
        wkhtmltopdf_path = r'C:\Program Files\wkhtmltopdf\bin\wkhtmltopdf.exe'

        # Replace placeholders and generate QR code
        html_content = replace_placeholders(html_content, recipient_email, random_values)

        if '{QR_CODE_PLACEHOLDER}' in html_content:
            qr_code_img = generate_qr_code(random.choice(config['links']), recipient_email, random_values, box_size=10, border=4)
            html_content = html_content.replace('{QR_CODE_PLACEHOLDER}', f'<img src="{qr_code_img}" alt="QR Code" style="width: 200px; height: 200px; dark=0044CC; light=F0F8FF;">')

        # Replace cid references with base64-encoded data URIs
        if image_paths:
            for cid, image_path in image_paths.items():
                if os.path.exists(image_path):
                    with open(image_path, "rb") as img_file:
                        encoded_image = b64encode(img_file.read()).decode('utf-8')
                    # Replace cid references in the HTML with the base64 data URI
                    html_content = html_content.replace(f'cid:{cid}', f'data:image/png;base64,{encoded_image}')
                else:
                    print(f"\033[91m[ERROR] Image file not found at path: {image_path}\033[0m")

        # Add CSS to control page layout
        html_content = f"""
        <html>
        <head>
            <style>
                body {{ font-size: 10pt; }}
                .page-break {{ page-break-after: always; }}
                img {{ max-width: 120%; }}
                table {{ width: 120%; }}
                td {{ padding: 5px; }}
            </style>
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """

        temp_html = f'temp_{uuid.uuid4().hex}.html'
        temp_pdf = f'temp_{uuid.uuid4().hex}.pdf'

        with open(temp_html, 'w', encoding='utf-8') as f:
            f.write(html_content)

        # Use wkhtmltopdf with options to control page size and margins
        command = [
            wkhtmltopdf_path,
            '--page-size', 'A4',
            '--margin-top', '10mm',
            '--margin-right', '10mm',
            '--margin-bottom', '10mm',
            '--margin-left', '10mm',
            '--enable-local-file-access',
            temp_html,
            temp_pdf
        ]
        result = subprocess.run(command, capture_output=True, text=True, check=True)

        with open(temp_pdf, 'rb') as pdf_file:
            pdf_content = pdf_file.read()

        return pdf_content
    except subprocess.CalledProcessError as e:
        print(f"\033[91m[ERROR] wkhtmltopdf failed: {e.stderr}\033[0m")
        return None
    except Exception as e:
        print(f"\033[91m[ERROR] Failed to convert HTML to PDF: {str(e)}\033[0m")
        return None
    finally:
        if os.path.exists(temp_html):
            os.remove(temp_html)
        if os.path.exists(temp_pdf):
            os.remove(temp_pdf)

def future_past_date(direction, num, unit):
    today = datetime.now()
    if direction == 'future':
        delta = timedelta(**{unit: int(num)})
        return (today + delta).strftime('%Y-%m-%d')
    elif direction == 'past':
        delta = timedelta(**{unit: int(num)})
        return (today - delta).strftime('%Y-%m-%d')
    return today.strftime('%Y-%m-%d')

def random_num(length):
    return ''.join(random.choices(string.digits, k=int(length)))

def random_str(length, case):
    chars = string.ascii_lowercase if case == 'lowercase' else string.ascii_uppercase
    return ''.join(random.choices(chars, k=int(length)))

def random_mix(length, case):
    chars = string.ascii_letters + string.digits
    if case == 'lowercase':
        chars = string.ascii_lowercase + string.digits
    elif case == 'uppercase':
        chars = string.ascii_uppercase + string.digits
    return ''.join(random.choices(chars, k=int(length)))


def get_favicon(domain):
    return f"https://logo.clearbit.com/{domain}"

def faker_data(locale, data_type, stability):
    fake = Faker(locale)
    Faker.seed(stability == 'stable')
    method = getattr(fake, data_type, None)
    return method() if method else f"Unsupported: {data_type}"


def generate_fake_company_email():
    company_domains = [
        'tech.com', 'systems.io', 'solutions.com', 'corp.net', 'enterprise.com',
        'global.com', 'innovations.net', 'group.com', 'industries.com', 'services.net',
        'consulting.com', 'software.io', 'analytics.com', 'partners.net', 'international.com'
    ]
    company_name = ''.join(random.choices(string.ascii_lowercase, k=random.randint(5, 10)))
    domain = random.choice(company_domains)
    return f"{company_name}@{domain}"

def generate_random_name():
    fake = Faker()
    return fake.name()

def mask_email(email, mask_char='*'):
    try:
        # Handle email addresses with multiple dots or special characters
        if '@' not in email:
            return email

        local_part, domain = email.split('@', 1)  # Split only at the first '@'

        # Mask local part
        if len(local_part) <= 2:
            masked_local = local_part
        else:
            masked_local = local_part[0] + mask_char * (len(local_part) - 2) + local_part[-1]

        # Handle domain with multiple dots
        domain_parts = domain.split('.')
        if len(domain_parts) < 2:
            return email  # Invalid domain format

        # Mask the domain name (everything before the last dot)
        domain_name = '.'.join(domain_parts[:-1])
        domain_ext = domain_parts[-1]

        if len(domain_name) <= 1:
            masked_domain = domain_name
        else:
            masked_domain = domain_name[0] + mask_char * (len(domain_name) - 1)

        return f"{masked_local}@{masked_domain}.{domain_ext}"
    except Exception as e:
        print(f"\033[93m[WARNING] Error masking email {email}: {str(e)}\033[0m")
        return email  # Return original email if masking fails


def replace_placeholders(text, recipient_email, random_values):
    email_prefix = recipient_email.split('@')[0]
    placeholders = {
        '{RECIPIENT_NAME}': email_prefix,
        '{RECIPIENT_EMAIL}': recipient_email,
        '{RECIPIENT_DOMAIN}': recipient_email.split('@')[1],
        '{RECIPIENT_DOMAIN_NAME}': recipient_email.split('@')[1].split('.')[0],
        '{RECIPIENT_$emailx}': mask_email(recipient_email),
        '{RECIPIENT_$emailv}': mask_email(recipient_email),
        '{CURRENT_DATE}': datetime.now().strftime('%Y-%m-%d'),
        '{CURRENT_TIME}': datetime.now().strftime('%H:%M:%S'),
        '{CURRENT_SECOND}': datetime.now().strftime('%S'),
        '{RANDOM_NUMBER10}': random_values['random_number10'],
        '{RANDOM_PHONENUMBER}': random_values['random_phonenumber'],
        '{RANDOM_STRING}': random_values['random_string'],
        '{RANDOM_MD5}': random_values['random_md5'],
        '{FAKE_COMPANY_EMAIL}': generate_fake_company_email(),
        '{FAKE_COMPANY_EMAIL_AND_FULLNAME}': f"{faker_data('en_US', 'name', 'random')} <{generate_fake_company_email()}>",
        '{RANDOM_PATH}': random_values['random_path'],
        '{RECIPIENT_BASE64_EMAIL}': base64.b64encode(recipient_email.encode()).decode(),
        '{RANDLINK}': random.choice(config['links']).replace("{RECIPIENT_EMAIL}", recipient_email),
        '{RAND-NAMES}': generate_random_name(),
        '{LINK}': random.choice(config['links']).replace("{RECIPIENT_EMAIL}", recipient_email),
        '{CURRENT_MINUTE}': datetime.now().strftime('%M'),
        '{CURRENT_HOUR}': datetime.now().strftime('%H'),
        '{CURRENT_DAY}': datetime.now().strftime('%d'),
        '{CURRENT_MONTH}': datetime.now().strftime('%m'),
        '{CURRENT_YEAR}': datetime.now().strftime('%Y'),
        '{CURRENT_DATE_LONG}': datetime.now().strftime('%B %d, %Y'),
        '{CURRENT_DATE_PLUS_TIME}': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        '{DOMAIN LOGO_FAVICON}': get_favicon(recipient_email.split('@')[1]),
        '{BASE64 ENCODED EMAIL_BASE64_EMAIL}': base64.b64encode(recipient_email.encode()).decode(),
    }

    for placeholder, value in placeholders.items():
        text = text.replace(placeholder, str(value))

    # Handle more complex placeholders
    text = re.sub(r'\{FUTURE_DATE_FUTURE_PAST\(\'(future|past)\',(\d+),\'(days|weeks)\'\)\}',
                  lambda m: future_past_date(m.group(1), m.group(2), m.group(3)), text)
    text = re.sub(r'\{SOME RANDOM NUMBER_RANDOM_NUM\((\d+)\)\}',
                  lambda m: random_num(m.group(1)), text)
    text = re.sub(r'\{SOME RANDOM TEXT_RANDOM_STR\((\d+),\'(lowercase|uppercase)\'\)\}',
                  lambda m: random_str(m.group(1), m.group(2)), text)
    text = re.sub(r'\{SOME RANDOM MIXED TEXT_RANDOM_MIX\((\d+),\'(lowercase|uppercase)\'\)\}',
                  lambda m: random_mix(m.group(1), m.group(2)), text)
    text = re.sub(r'\{CRYPTO RANDOM MD5_RANDOM_MD5\(\'(lowercase|uppercase)\'\)\}',
                  lambda m: random_values['random_md5'].lower() if m.group(1) == 'lowercase' else random_values['random_md5'].upper(), text)
    text = re.sub(r'\{BASE64 ENCODED_BASE64_ENCODE\(\'(.+?)\'\)\}',
                  lambda m: base64.b64encode(m.group(1).encode()).decode(), text)
    text = re.sub(r'\{BASE64 DECODED_BASE64_DECODE\(\'(.+?)\'\)\}',
                  lambda m: base64.b64decode(m.group(1)).decode(), text)
    text = re.sub(r'\{UPPERCASE_UPPERCASE\(\'(.+?)\'\)\}',
                  lambda m: m.group(1).upper(), text)
    text = re.sub(r'\{LOWERCASE_LOWERCASE\(\'(.+?)\'\)\}',
                  lambda m: m.group(1).lower(), text)
    text = re.sub(r'\{CAPITALIZE_CAPITALIZE\(\'(.+?)\'\)\}',
                  lambda m: m.group(1).capitalize(), text)
    text = re.sub(r'\{NAMECASE_NAMECASE\(\'(.+?)\'\)\}',
                  lambda m: ' '.join(word.capitalize() for word in m.group(1).split()), text)
    text = re.sub(r'\{SENTENCECASE_SENTENCECASE\(\'(.+?)\'\)\}',
                  lambda m: m.group(1).capitalize(), text)
    text = re.sub(r'\{A LIST OF FAKE GENERATED INFORMATIONS_FAKER\(\'(.+?)\',\'(.+?)\',\'(.+?)\'\)\}',
                  lambda m: faker_data(m.group(1), m.group(2), m.group(3)), text)

    # Handle QR code placeholder
    if '{QR CODE IMAGE_QRCODE_URL}' in text:
        link = random.choice(config['links'])
        link = replace_placeholders(link, recipient_email, random_values)
        qr_code_image = generate_qr_code(link)
        text = text.replace('{QR CODE IMAGE_QRCODE_URL}', f'<img src="{qr_code_image}" alt="QR Code">')

    # Handle EMBEDDED FILE placeholder
    if '{EMBEDDED FILE_EMBEDDED_URL}' in text:
        # Implement the logic for embedding a file here
        # This might involve reading a file and encoding it, or generating a URL
        pass

    return text

def wait_for_user_exit():
    """
    Keeps the program running until the user decides to close it.
    """
    print("\n Job well done. Press Ctrl+C to exit or close the terminal.")
    try:
        while True:
            time.sleep(1)
    except KeyboardInterrupt:
        print("\nApplication closed. Goodbye!")

def main():
    # Ask user to enter License key
    pc_identifier = get_pc_identifier()
    license_key = load_license()

    # If no license key is stored, prompt the user to enter one
    if not license_key:
        license_key = prompt_license_key(pc_identifier)
    else:
        # Validate the loaded license key
        if not validate_license(license_key, pc_identifier):
            print("Stored license key is invalid. Please enter a valid license key.")
            license_key = prompt_license_key(pc_identifier)

    # Start a separate thread to check the license continuously
    threading.Thread(target=check_license_continuously, args=(license_key, pc_identifier), daemon=True).start()

    # Check if user has license key

    print_banner()
    print("\033[92m[INFO] Starting email script\033[0m")

    # Load email list from config
    email_list_path = os.path.join(base_path, 'list.txt')
    with open(email_list_path, 'r') as list_file:
        email_list = list_file.readlines()

    total_emails = len(email_list)
    print(f"\033[94m[INFO] Loaded {total_emails} email addresses from list.txt\033[0m")

    # Send emails concurrently based on the selected method from config
    successful_sends, failed_sends = send_emails_concurrently(email_list, body="Test email body")

    print_summary(successful_sends, failed_sends, total_emails)

    # Tell user to close App
    wait_for_user_exit()


if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\n\033[93m[INFO] Script execution was interrupted by user.\033[0m")
        sys.exit(0)
